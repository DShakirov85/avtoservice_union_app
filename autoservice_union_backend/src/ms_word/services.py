from docx import Document

from members.schemas import   (
    CompanyLeaderForm,
    IndividualEntrepreneurForm,
    LegalEntityForm
)

from members.schemas import BankDetails
from logger import logger


class MSWordManager:
    def create_leader_participation_form(
            self,
            values: CompanyLeaderForm
    ):

        document = Document("templates/leader_participation_form_new.docx")
        bank_details = values.bank_details
        if not bank_details:
            bank_details = BankDetails(
                bank_name="",
                account_number="",
                correspondent_account="",
                bik=""
            )
        replace_values = {
            "{{full_name}}": values.full_name,
            "{{passport}}": values.passport_data,
            "{{legal_address}}": values.legal_address,
            "{{inn}}": values.company_inn,
            "{{company_name}}": values.company_name,
            "{{okved_codes}}": values.okved_codes,
            "{{actual_address}}": values.actual_address,
            "{{mailing_address}}": values.mailing_address,
            "{{company_phone}}": values.company_phone,
            "{{email}}": values.email,
            "{{personal_phone}}": values.personal_phone,
            "{{website}}": values.website,
        }
        for placeholder, value in replace_values.items():
            self.replace_placeholder(document, placeholder, value)

        document.save(f"files/leader_participation_form_{values.company_name}.docx")
        return f"files/leader_participation_form_{values.company_name}.docx"

    def create_individual_entrepreneur_participation_form(
            self,
            values: IndividualEntrepreneurForm
    ):
        bank_details = values.bank_details
        if not bank_details:
            bank_details = BankDetails(
                bank_name="",
                account_number="",
                correspondent_account="",
                bik=""
            )
        document = Document("templates/individual_participation_form_new.docx")
        replace_values = {
            "{{full_name}}": values.full_name,
            "{{registration_address}}": values.registration_address,
            "{{tax_registration_date}}": values.tax_registration_date,
            "{{company_inn}}": values.company_inn,
            "{{ogrnip}}": values.ogrnip,
            "{{okved_codes}}": values.okved_codes,
            "{{actual_address}}": values.actual_address,
            "{{mailing_address}}": values.mailing_address,
            "{{company_phone}}": values.company_phone or "",
            "{{personal_phone}}": values.personal_phone,
            "{{email}}": values.email or "",
            "{{website}}": values.website,
            "{{bank}}": bank_details.bank_name,
            "{{account_number}}": bank_details.account_number,
            "{{correspondent_account}}": bank_details.correspondent_account,
            "{{bik}}": bank_details.bik,
        }
        for placeholder, value in replace_values.items():
            self.replace_placeholder(document, placeholder, value)

        document.save(f"files/individual_participation_form_{values.full_name}.docx")
        return f"files/individual_participation_form_{values.full_name}.docx"

    def create_legal_entity_participation_form(
        self,
        values: LegalEntityForm
    ):
        document = Document("templates/legal_entity_participation_form_new.docx")
        bank_details = values.bank_details
        if not bank_details:
            bank_details = BankDetails(
                bank_name="",
                account_number="",
                correspondent_account="",
                bik=""
            )
        replace_values = {
            "{{full_name}}": values.full_name,
            "{{company_name}}": values.company_name,
            "{{company_kpp}}": values.company_kpp,
            "{{company_ogrn}}": values.company_ogrn,
            "{{director_position}}": values.director_position,
            "{{director_full_name}}": values.director_full_name,
            "{{registration_address}}": values.registration_address,
            "{{registration_date}}": values.registration_date.strftime("%Y-%m-%d"),
            "{{company_inn}}": values.company_inn,
            "{{ogrnip}}": values.ogrnip,
            "{{okved_codes}}": values.okved_codes,
            "{{actual_address}}": values.actual_address,
            "{{legal_address}}": values.legal_address,
            "{{mailing_address}}": values.mailing_address,
            "{{company_phone}}": values.company_phone,
            "{{personal_phone}}": values.personal_phone,
            "{{email}}": values.email,
            "{{website}}": values.website,
            "{{bank_name}}": bank_details.bank_name,
            "{{account_number}}": bank_details.account_number,
            "{{correspondent_account}}": bank_details.correspondent_account,
            "{{bik}}": bank_details.bik,
        }
        for placeholder, value in replace_values.items():
            self.replace_placeholder(document, placeholder, value)

        document.save(f"files/legal_entity_participation_form_{values.full_name}.docx")
        return f"files/legal_entity_participation_form_{values.full_name}.docx"


    def create_survey_form(
            self,
            values: CompanyLeaderForm | IndividualEntrepreneurForm | LegalEntityForm
    ):
        document = Document("templates/anketa_new.docx")
        bank_details = values.bank_details
        if not bank_details:
            bank_details = BankDetails(
                bank_name="",
                account_number="",
                correspondent_account="",
                bik=""
            )
        try:
            company_name = values.company_name
            company_short_name = values.company_short_name
        except AttributeError:
            company_name = f"ИП {values.full_name}"
            company_short_name = f"ИП {values.full_name}"
        social_media = values.social_media

        replace_values = {
            "{{full_name}}": values.full_name,
            "{{company_name}}": company_name,
            "{{company_short_name}}": company_short_name,
            "{{registration_date}}": values.registration_date,
            "{{legal_address}}": values.legal_address,
            "{{actual_address}}": values.actual_address,
            "{{mailing_address}}": values.mailing_address,
            "{{company_inn}}": values.company_inn,
            "{{okved_codes}}": values.okved_codes,
            "{{bank}}": bank_details.bank_name,
            "{{account_number}}": bank_details.account_number,
            "{{correspondent_account}}": bank_details.correspondent_account,
            "{{bik}}": bank_details.bik,
            "{{locations}}": values.locations,
            "{{employee_count}}": values.employee_count,
            "{{service_points_count}}": values.service_points_count,
            "{{service_zones_photos}}": values.service_zones_photos,
            "{{facilities}}": values.facilities,
            "{{specialization}}": values.specialization,
            "{{website}}": values.website,
            "{{social_media}}": social_media,
            "{{interested_services}}": values.interested_services,
            "{{can_help_with}}": values.can_help_with,
            "{{recommended_partners}}": values.recommended_partners,
            "{{submission_date}}": values.submission_date,
        }
        for placeholder, value in replace_values.items():
            self.replace_placeholder(document, placeholder, value)
        document.save(f"files/survey_form_{values.full_name}.docx")
        return f"files/survey_form_{values.full_name}.docx"

    def create_principle(
        self,
        values: CompanyLeaderForm | IndividualEntrepreneurForm | LegalEntityForm
    ):
        document = Document("templates/principles.docx")
        replace_values = {
            "{{full_name}}": values.full_name,
         }
        for placeholder, value in replace_values.items():
            self.replace_placeholder(document, placeholder, value)
        document.save(f"files/principles_{values.full_name}.docx")
        return f"files/principles_{values.full_name}.docx"

    def create_personal_data_consent(
        self,
        values: CompanyLeaderForm | IndividualEntrepreneurForm | LegalEntityForm
    ):
        document = Document("templates/personal_data_consent.docx")
        replace_values = {
            "{{full_name}}": values.full_name,
         }
        for placeholder, value in replace_values.items():
            self.replace_placeholder(document, placeholder, value)
        document.save(f"files/personal_data_consent_{values.full_name}.docx")
        return f"files/personal_data_consent_{values.full_name}.docx"


    def replace_placeholder(
            self,
            document: Document,
            placeholder: str,
            value: str
    ):
        for paragraph in document.paragraphs:
            if placeholder in paragraph.text:
                try:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                except TypeError:
                    paragraph.text = paragraph.text.replace(placeholder, "")

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        try:
                            cell.text = cell.text.replace(placeholder, value)
                        except TypeError:
                            cell.text = cell.text.replace(placeholder, "")
